# -*- coding: utf-8 -*-
import pandas as pd
import numpy as np
from datetime import datetime
from docx import Document
from docx.shared import Pt
import argparse
import os
import re

# ---------- 列名常量 ----------
COL_DATE = '监测时间（年/月/日）'
COL_VALUE = '监测指标值'
COL_TYPE = '防控区类型'
COL_REGION = '地市-区/县/市-街道/乡/镇'
COL_COMMUNITY = '社区/村居'
COL_ADDR1 = '监测地址（地图定位版）'
COL_ADDR2 = '监测地址（如“监测地址”定位字段不可用，可手填；如定位可用，不需要重复填写）'
COL_DAYS = '距末例天数（自动计算）'
COL_METHOD = '监测方法（如BI/RI/MOI/ADI等）'
COL_CITY = '地市'
COL_DISTRICT = '区/县/市'
COL_STREET = '街道/乡/镇'

# 地市排序顺序
CITY_ORDER = [
    '广州市', '深圳市', '珠海市', '汕头市', '佛山市', '韶关市', '河源市',
    '梅州市', '惠州市', '汕尾市', '东莞市', '中山市', '江门市', '阳江市',
    '湛江市', '茂名市', '肇庆市', '清远市', '潮州市', '揭阳市', '云浮市'
]

def parse_date(date_str):
    for fmt in ('%Y-%m-%d', '%Y/%m/%d', '%Y%m%d'):
        try:
            return datetime.strptime(date_str, fmt)
        except Exception:
            continue
    raise ValueError(f'无法解析日期: {date_str}')

def load_and_filter(excel_path, target_date, exclude_str):
    """加载Excel并执行第1-5步筛选。返回 (original_df, filtered_df)."""
    original = pd.read_excel(excel_path, sheet_name=0, dtype=str)
    df = original.copy()

    # 1. 筛选日期
    df[COL_DATE] = pd.to_datetime(df.get(COL_DATE), errors='coerce')
    target = pd.to_datetime(target_date)
    df = df[df[COL_DATE].dt.date == target.date()].copy()

    # 2. 去除监测指标值为空
    df[COL_VALUE] = pd.to_numeric(df.get(COL_VALUE), errors='coerce')
    df = df[df[COL_VALUE].notna()].copy()

    # 3. 筛选防控区类型
    df = df[df.get(COL_TYPE).isin(['核心区', '警戒区'])].copy()

    # 4. 排除包含指定字段的行
    if exclude_str and exclude_str.strip():
        df = df[~df.get(COL_REGION, '').str.contains(exclude_str, na=False)].copy()

    # 5. 距末例天数条件
    df[COL_DAYS] = pd.to_numeric(df.get(COL_DAYS), errors='coerce')
    df = df[(df[COL_DAYS] <= 5) | (df[COL_DAYS] > 40000)].copy()

    return original, df

def process_bi_ssi(df):
    """处理BI和SSI数据（第6-9步），返回处理后的BI表"""
    bi_ssi = df[df.get(COL_METHOD).isin(['布雷图指数BI', '标准间指数SSI'])].copy()
    if bi_ssi.empty:
        return pd.DataFrame(columns=df.columns)

    # 分离BI和SSI
    bi = bi_ssi[bi_ssi[COL_METHOD] == '布雷图指数BI'].copy()
    bi[COL_VALUE] = pd.to_numeric(bi[COL_VALUE], errors='coerce')
    bi['_src'] = 'BI'

    ssi = bi_ssi[bi_ssi[COL_METHOD] == '标准间指数SSI'].copy()
    ssi[COL_VALUE] = pd.to_numeric(ssi[COL_VALUE], errors='coerce') * 2
    ssi[COL_METHOD] = '布雷图指数BI'
    ssi['_src'] = 'SSI'

    combined = pd.concat([bi, ssi], ignore_index=True)

    # 关键字段分组，填充空值以避免分组错误
    key_cols = [COL_REGION, COL_COMMUNITY, COL_ADDR1, COL_ADDR2, COL_TYPE]
    for col in key_cols:
        combined[col] = combined[col].fillna('')

    def agg_func(group):
        ssi_rows = group[group['_src'] == 'SSI']
        bi_rows = group[group['_src'] == 'BI']
        bi_max = bi_rows[COL_VALUE].max() if not bi_rows.empty else np.nan
        ssi_max = ssi_rows[COL_VALUE].max() if not ssi_rows.empty else np.nan
        # 新逻辑：若SSI转换后>5，则取max(BI, SSI)，否则只取BI
        if not np.isnan(ssi_max) and ssi_max > 5:
            if np.isnan(bi_max):
                final_value = ssi_max
            else:
                final_value = max(bi_max, ssi_max)
        else:
            final_value = bi_max
        first_row = group.iloc[0].copy()
        first_row[COL_VALUE] = final_value
        return first_row

    grouped = combined.groupby(key_cols, as_index=False)
    bi_final = grouped.apply(lambda g: agg_func(g)).reset_index(drop=True)
    bi_final.drop(columns=['_src'], inplace=True, errors='ignore')
    return bi_final

def process_adi(df):
    """处理ADI数据（第10步）"""
    adi = df[df.get(COL_METHOD) == '成蚊密度指数法ADI'].copy()
    if adi.empty:
        return pd.DataFrame(columns=df.columns)
    adi[COL_VALUE] = pd.to_numeric(adi[COL_VALUE], errors='coerce')
    adi = adi.dropna(subset=[COL_VALUE]).copy()
    if adi.empty:
        return pd.DataFrame(columns=df.columns)

    key_cols = [COL_REGION, COL_COMMUNITY, COL_ADDR1, COL_ADDR2, COL_TYPE]
    for col in key_cols:
        adi[col] = adi[col].fillna('')

    # 对每组取最大值的那一行
    idx = adi.groupby(key_cols, as_index=False)[COL_VALUE].idxmax()
    # idx may contain NaN if group empty, filter valid ints
    idx = [i for i in idx if pd.notna(i)]
    if not idx:
        return pd.DataFrame(columns=df.columns)
    adi_final = adi.loc[idx].reset_index(drop=True)
    return adi_final

def split_region(region_str):
    """尝试从区域字段中提取 (city, district, street)。
    对常见格式做更稳健的处理：去掉省名后，用正则按市/区/县/镇等后缀提取。
    """
    if not isinstance(region_str, str):
        return '', '', ''
    s = region_str.strip()
    # 尝试去掉省名前缀（如广东省）
    s = re.sub(r'^.*?省', '', s)
    # 先尝试匹配市/区/县/街道等
    m = re.match(r'(?P<city>.*?市)?(?P<district>.*?(?:区|县|市))?(?P<street>.*)', s)
    if not m:
        return '', '', ''
    city = (m.group('city') or '').strip()
    district = (m.group('district') or '').strip()
    street = (m.group('street') or '').strip()
    return city, district, street

def risk_level(value):
    """返回 (label, hex_color)；hex_color 为 6 位大写字符串，不带#。"""
    try:
        v = float(value)
    except Exception:
        return '未知', 'FFFFFF'
    if v < 5:
        return '安全', '00FF00'  # 绿色
    elif v < 10:
        return '低风险', 'FFFF00'  # 黄色
    elif v < 20:
        return '中风险', 'FFA500'  # 橘黄
    else:
        return '高风险', 'FF0000'  # 红色

def generate_report_doc(bi_df, adi_df, date_str, exclude_str, output_path):
    doc = Document()
    month_day = date_str[5:7] + '月' + date_str[8:10] + '日'
    title = f'省媒介伊蚊传染病疫情蚊媒监测情况（{month_day} 20：00）'
    doc.add_heading(title, 0)

    # 统计
    total = len(bi_df)
    core = len(bi_df[bi_df.get(COL_TYPE) == '核心区'])
    alert = len(bi_df[bi_df.get(COL_TYPE) == '警戒区'])
    # 风险分级
    high = bi_df[bi_df.get(COL_VALUE) >= 20]
    mid = bi_df[(bi_df.get(COL_VALUE) >= 10) & (bi_df.get(COL_VALUE) < 20)]
    low = bi_df[(bi_df.get(COL_VALUE) >= 5) & (bi_df.get(COL_VALUE) < 10)]
    safe = bi_df[bi_df.get(COL_VALUE) < 5]

    # 获取所有地市（从bi_df）
    all_cities = set()
    for _, row in bi_df.iterrows():
        city, _, _ = split_region(row.get(COL_REGION, '') or '')
        if city:
            all_cities.add(city)
    sorted_cities = [c for c in CITY_ORDER if c in all_cities]
    # 添加未在 CITY_ORDER 的城市（按字典序）
    remaining = sorted(list(all_cities - set(sorted_cities)))
    sorted_cities += remaining

    # 构建地市字符串
    if len(sorted_cities) == 0:
        city_text = '全省'
    elif len(sorted_cities) == 1:
        city_text = sorted_cities[0]
    else:
        city_text = '、'.join([c.replace('市', '') for c in sorted_cities[:-1]]) + f'和{sorted_cities[-1]}'

    # 处理排除备注
    if exclude_str:
        city_text += f'（{exclude_str}除外）'

    p = doc.add_paragraph()
    p.add_run(f'{city_text}报告媒介伊蚊监测点共{total}个，核心区{core}个，警戒区{alert}个，高风险区{len(high)}个，中风险区{len(mid)}个，低风险区{len(low)}个，安全区{len(safe)}个。')

    # 分段描述各风险区
    def describe_risk(df_risk, risk_name):
        if df_risk.empty:
            return
        doc.add_paragraph(f'{risk_name}风险区：')
        # 按值升序排序
        df_sorted = df_risk.sort_values(COL_VALUE)
        for _, row in df_sorted.iterrows():
            city, district, street = split_region(row.get(COL_REGION, '') or '')
            # 处理东莞、中山去除“市辖区”
            if city in ['东莞市', '中山市']:
                district = district.replace('市辖区', '')
            loc = f'{city}（{district}-{street}）' if district and street else (city or '未知')
            text = f'{loc}：{row.get(COL_COMMUNITY, "")}（{row.get(COL_VALUE, "")}）'
            doc.add_paragraph(text, style='List Bullet')

    describe_risk(high, '高')
    describe_risk(mid, '中')
    describe_risk(low, '低')
    # 安全区不逐一列出，可省略

    doc.save(output_path)

def generate_table_doc(bi_df, adi_df, date_str, output_path):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    doc = Document()
    month_day = date_str[5:7] + '月' + date_str[8:10] + '日'
    title = f'全省媒介伊蚊传染病疫点重点镇（街道）蚊媒密度监测村居一览表（{month_day} 20：00）'
    doc.add_heading(title, 0)

    # 处理BI表格
    doc.add_heading('一、布雷图指数（BI）监测', level=1)
    if not bi_df.empty:
        # 准备数据
        data = []
        for _, row in bi_df.iterrows():
            city, district, street = split_region(row.get(COL_REGION, '') or '')
            if city in ['东莞市', '中山市']:
                district = district.replace('市辖区', '')
            district_clean = re.sub(r'[市区县]', '', district or '')
            community = row.get(COL_COMMUNITY, '') or ''
            value = row.get(COL_VALUE, '')
            risk, hex_color = risk_level(value)
            data.append({
                'city': city or '',
                'district': district_clean,
                'street': street or '',
                'community': community,
                'value': value,
                'risk': risk,
                'color': hex_color,
                'type': row.get(COL_TYPE, '')
            })
        df_display = pd.DataFrame(data)
        # 排序
        df_display['city_rank'] = df_display['city'].map(lambda x: CITY_ORDER.index(x) if x in CITY_ORDER else 999)
        df_display = df_display.sort_values(['city_rank', 'district', 'street', 'community', 'type'])

        # 创建表格
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        # 标题行
        hdr = table.rows[0].cells
        hdr[0].text = '地市'
        hdr[1].text = '区/县'
        hdr[2].text = '街道/乡/镇'
        hdr[3].text = '村居'
        hdr[4].text = 'BI'
        hdr[5].text = '风险水平*'

        last_city = None
        for _, row in df_display.iterrows():
            cells = table.add_row().cells
            # 合并相同地市显示（这里只是空文本避免重复显示）
            if row['city'] != last_city:
                cells[0].text = row['city']
                last_city = row['city']
            else:
                cells[0].text = ''
            cells[1].text = str(row['district'])
            cells[2].text = str(row['street'])
            cells[3].text = str(row['community'])
            cells[4].text = str(row['value'])
            cells[5].text = str(row['risk'])

            # 设置单元格底纹颜色（用 hex）
            tc = cells[5]._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            hex_color = row['color'] if isinstance(row['color'], str) else 'FFFFFF'
            # Ensure 6-char uppercase
            hex_color = (hex_color.replace('#', '').upper() + 'FFFFFF')[:6]
            shd.set(qn('w:fill'), hex_color)
            tcPr.append(shd)

        # 尝试设置首行作为表头重复（文档兼容性可能因 docx 版本不同）
        # note: exact behavior may vary across Word versions
        try:
            tr = table.rows[0]._tr
            tblPr = table._tbl.get_or_add_tblPr()
            tblPr.append(OxmlElement('w:tblHeader'))
        except Exception:
            pass

    doc.add_page_break()

    # 处理ADI表格
    doc.add_heading('二、成蚊密度（ADI）监测', level=1)
    if not adi_df.empty:
        data = []
        for _, row in adi_df.iterrows():
            city, district, street = split_region(row.get(COL_REGION, '') or '')
            if city in ['东莞市', '中山市']:
                district = district.replace('市辖区', '')
            district_clean = re.sub(r'[市区县]', '', district or '')
            data.append({
                'city': city or '',
                'district': district_clean,
                'street': street or '',
                'community': row.get(COL_COMMUNITY, '') or '',
                'value': row.get(COL_VALUE, ''),
                'type': row.get(COL_TYPE, '')
            })
        df_display = pd.DataFrame(data)
        df_display['city_rank'] = df_display['city'].map(lambda x: CITY_ORDER.index(x) if x in CITY_ORDER else 999)
        df_display = df_display.sort_values(['city_rank', 'district', 'street', 'community', 'type'])

        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        hdr[0].text = '地市'
        hdr[1].text = '区/县'
        hdr[2].text = '街道/乡/镇'
        hdr[3].text = '村居'
        hdr[4].text = 'ADI'

        last_city = None
        for _, row in df_display.iterrows():
            cells = table.add_row().cells
            if row['city'] != last_city:
                cells[0].text = row['city']
                last_city = row['city']
            else:
                cells[0].text = ''
            cells[1].text = str(row['district'])
            cells[2].text = str(row['street'])
            cells[3].text = str(row['community'])
            cells[4].text = str(row['value'])

        try:
            tr = table.rows[0]._tr
            tblPr = table._tbl.get_or_add_tblPr()
            tblPr.append(OxmlElement('w:tblHeader'))
        except Exception:
            pass

    doc.save(output_path)

def generate_deleted_log(original_df, filtered_df, exclude_str, output_path):
    doc = Document()
    doc.add_heading('被删除数据情况说明', 0)
    deleted_count = len(original_df) - len(filtered_df)
    doc.add_paragraph(f'共删除 {deleted_count} 条记录。')
    if exclude_str:
        doc.add_paragraph(f'原因示例：排除包含字段 "{exclude_str}" 的记录。')
    # 可扩展为写出每条被删记录的详细信息（行号/原因）
    doc.save(output_path)

def main():
    parser = argparse.ArgumentParser()
    parser.add_argument('--excel', required=True, help='Excel文件路径')
    parser.add_argument('--date', required=True, help='监测日期，格式YYYY-MM-DD')
    parser.add_argument('--exclude', default='', help='需要排除的行政区字段，如"荔湾区"')
    parser.add_argument('--output', default='./output', help='输出目录')
    args = parser.parse_args()

    os.makedirs(args.output, exist_ok=True)

    original_df, df = load_and_filter(args.excel, args.date, args.exclude)
    bi_df = process_bi_ssi(df)
    adi_df = process_adi(df)

    date_str = args.date
    exclude_str = args.exclude
    report_path = os.path.join(args.output, f'省媒介伊蚊传染病疫情蚊媒监测情况（{date_str[5:7]}月{date_str[8:10]}日 20：00）.docx')
    table_path = os.path.join(args.output, f'全省媒介伊蚊传染病疫点重点镇（街道）蚊媒密度监测村居一览表（{date_str[5:7]}月{date_str[8:10]}日 20：00）.docx')
    deleted_log_path = os.path.join(args.output, '被删除数据情况说明.docx')

    generate_report_doc(bi_df, adi_df, date_str, exclude_str, report_path)
    generate_table_doc(bi_df, adi_df, date_str, table_path)
    generate_deleted_log(original_df, df, exclude_str, deleted_log_path)

    print('处理完成！')

if __name__ == '__main__':
    main()
